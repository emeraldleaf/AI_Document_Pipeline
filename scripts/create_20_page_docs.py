#!/usr/bin/env python3
"""
Create 20-page test documents (PDF and Word) for testing multi-page document processing.
"""

from pathlib import Path
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet

def create_20_page_word_doc():
    """Create a 20-page Word document."""
    doc = Document()

    # Add title page
    doc.add_heading('Annual Business Report 2025', 0)
    doc.add_heading('Company Performance Analysis', level=1)
    doc.add_paragraph('Prepared by: Finance Department')
    doc.add_paragraph('Date: November 4, 2025')
    doc.add_page_break()

    # Add 19 more pages with different content
    sections = [
        ("Executive Summary", "This comprehensive report provides detailed insights into our company's performance throughout 2025. Key metrics show significant growth across all departments with revenue increasing by 45% year-over-year."),
        ("Q1 Performance", "First quarter showed exceptional results with revenue of $2.5M. New client acquisition increased by 30%. Team expansion added 15 new employees across engineering and sales."),
        ("Q2 Performance", "Second quarter maintained momentum with $2.8M revenue. Product launches in April and May contributed significantly. Customer satisfaction scores reached 92%."),
        ("Q3 Performance", "Third quarter revenue reached $3.1M. International expansion began with new offices in London and Tokyo. Partnership agreements signed with three major enterprise clients."),
        ("Q4 Performance", "Fourth quarter projections indicate $3.5M revenue. Holiday season promotional campaigns showing strong results. Year-end targets on track to exceed expectations."),
        ("Sales Analysis", "Sales team exceeded all quarterly targets. Top performing regions include North America (40% growth) and Europe (35% growth). New CRM implementation improved conversion rates by 25%."),
        ("Marketing Initiatives", "Marketing campaigns generated 15,000 new leads. Social media engagement increased 120%. Content marketing strategy produced 50+ high-quality articles and case studies."),
        ("Product Development", "Engineering team delivered 12 major feature releases. Customer-requested features implemented based on feedback surveys. Technical debt reduced by 30% through refactoring initiatives."),
        ("Operations Review", "Operational efficiency improved with new automation tools. Customer support response time reduced from 24 hours to 6 hours. Infrastructure costs decreased 20% through cloud optimization."),
        ("Financial Highlights", "Gross profit margin increased to 68%. Operating expenses maintained at 45% of revenue. Cash reserves grew to $8M providing strong runway for 2026 expansion plans."),
        ("Employee Engagement", "Employee satisfaction survey results at all-time high of 87%. New benefits program launched including remote work options. Training programs completed by 95% of staff."),
        ("Technology Stack", "Migration to microservices architecture completed successfully. API performance improved 3x through optimization. Security audit passed with zero critical vulnerabilities."),
        ("Customer Success", "Net Promoter Score increased from 45 to 72. Customer retention rate at 94%. Premium tier upgrades increased by 40% indicating strong product value proposition."),
        ("Market Position", "Market share grew from 8% to 12% in primary vertical. Brand recognition studies show 35% improvement. Competitive analysis indicates strong positioning against top three competitors."),
        ("Risk Assessment", "Identified key risks include supply chain dependencies and regulatory changes. Mitigation strategies implemented including vendor diversification. Compliance team expanded to address new regulations."),
        ("Strategic Initiatives", "Five-year strategic plan approved by board. Focus areas include AI/ML integration, global expansion, and enterprise partnerships. Investment budget allocated for 2026 initiatives."),
        ("Sustainability Report", "Carbon footprint reduced 25% through remote work policies. Renewable energy sources now power 60% of operations. Sustainability goals align with 2030 climate commitments."),
        ("Community Impact", "Charitable contributions totaled $250K supporting education and technology access. Employee volunteer program engaged 80% of workforce. Scholarship program launched for underserved communities."),
        ("2026 Outlook", "Conservative projections estimate 40% revenue growth to $14M. Hiring plan includes 30 new positions. Product roadmap includes three major releases and expansion into two new market segments.")
    ]

    for i, (title, content) in enumerate(sections, 2):
        doc.add_heading(f'Chapter {i-1}: {title}', level=1)
        doc.add_paragraph(content)
        doc.add_paragraph('')
        doc.add_paragraph(f'Additional analysis and detailed metrics for {title.lower()} demonstrate continued positive trends. Strategic focus remains on sustainable growth while maintaining operational excellence.')
        doc.add_paragraph('')
        doc.add_paragraph(f'Key takeaway: {title} metrics indicate strong performance and positive trajectory for upcoming quarters. Team collaboration and execution excellence remain critical success factors.')

        if i < 20:  # Don't add page break after last page
            doc.add_page_break()

    # Save the document
    output_path = Path("documents/input/annual_report_20_pages.docx")
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"✓ Created 20-page Word document: {output_path}")
    return output_path

def create_20_page_pdf():
    """Create a 20-page PDF document."""
    output_path = Path("documents/input/technical_manual_20_pages.pdf")
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = SimpleDocTemplate(str(output_path), pagesize=letter)
    styles = getSampleStyleSheet()
    story = []

    # Title page
    story.append(Paragraph("Technical Implementation Manual", styles['Title']))
    story.append(Spacer(1, 12))
    story.append(Paragraph("System Architecture and Deployment Guide", styles['Heading1']))
    story.append(Spacer(1, 12))
    story.append(Paragraph("Version 2.0 - November 2025", styles['Normal']))
    story.append(PageBreak())

    # Add 19 more pages with technical content
    chapters = [
        ("Introduction", "This technical manual provides comprehensive documentation for system architecture, deployment procedures, and maintenance guidelines. The system utilizes microservices architecture with containerized deployments."),
        ("System Requirements", "Minimum requirements include: 16GB RAM, 4 CPU cores, 100GB storage. Recommended configuration: 32GB RAM, 8 CPU cores, 500GB SSD storage. Network bandwidth minimum 100Mbps for optimal performance."),
        ("Architecture Overview", "The system follows a three-tier architecture pattern with presentation layer (React frontend), application layer (Node.js APIs), and data layer (PostgreSQL with Redis caching). Load balancing via Nginx."),
        ("Database Configuration", "PostgreSQL 15+ required with pgvector extension for vector similarity search. Connection pooling configured with max 50 connections. Automated backups scheduled daily at 2 AM UTC."),
        ("API Endpoints", "REST API provides endpoints for document upload, classification, search, and retrieval. Authentication via JWT tokens with 24-hour expiration. Rate limiting at 1000 requests per hour per user."),
        ("Authentication Setup", "OAuth2 authentication flow with support for Google and Microsoft identity providers. Role-based access control (RBAC) with admin, user, and readonly roles. Session management via Redis."),
        ("Deployment Guide", "Docker Compose deployment for development. Kubernetes manifests provided for production. Environment variables configured via .env file. Secrets managed through HashiCorp Vault integration."),
        ("Monitoring Setup", "Prometheus metrics collection with Grafana dashboards. Custom metrics include request latency, classification accuracy, and database query performance. Alert rules configured for critical thresholds."),
        ("Logging Configuration", "Structured logging using JSON format. Log levels: DEBUG, INFO, WARNING, ERROR. Logs aggregated via ELK stack (Elasticsearch, Logstash, Kibana). Retention period 90 days."),
        ("Security Considerations", "TLS 1.3 encryption for all communications. SQL injection prevention via parameterized queries. XSS protection with Content Security Policy headers. Regular security audits scheduled quarterly."),
        ("Performance Tuning", "Database query optimization with proper indexing strategies. API response caching for frequently accessed data. CDN integration for static assets. Connection pooling and keep-alive configurations."),
        ("Backup Procedures", "Automated daily backups with 30-day retention. Point-in-time recovery capability. Backup verification tests performed weekly. Disaster recovery plan documented with RTO of 4 hours."),
        ("Scaling Strategy", "Horizontal scaling via container orchestration. Auto-scaling rules based on CPU and memory thresholds. Database read replicas for query distribution. CDN for geographic distribution."),
        ("Testing Procedures", "Unit tests with 85% code coverage requirement. Integration tests for API endpoints. End-to-end tests using Playwright. Performance tests with JMeter simulating 1000 concurrent users."),
        ("Troubleshooting Guide", "Common issues include connection timeouts (check network), high memory usage (review query patterns), slow classification (verify Ollama service). Debug mode enabled via DEBUG=true environment variable."),
        ("Maintenance Schedule", "Weekly security updates. Monthly dependency updates. Quarterly major version upgrades. Annual architecture review. Database maintenance window: Sunday 2-4 AM UTC."),
        ("Migration Procedures", "Database migration using Alembic. Zero-downtime deployments via blue-green strategy. Rollback procedures documented for each release. Data migration scripts version controlled."),
        ("API Documentation", "OpenAPI 3.0 specification available at /docs endpoint. Interactive API testing via Swagger UI. Code examples provided in Python, JavaScript, and cURL. Webhook integration guide included."),
        ("Appendix", "Additional resources include configuration templates, sample data sets, and integration examples. Community support available via GitHub discussions. Enterprise support contracts available for production deployments.")
    ]

    for i, (title, content) in enumerate(chapters, 2):
        story.append(Paragraph(f"Chapter {i-1}: {title}", styles['Heading1']))
        story.append(Spacer(1, 12))
        story.append(Paragraph(content, styles['Normal']))
        story.append(Spacer(1, 12))
        story.append(Paragraph(f"Detailed implementation notes: The {title.lower()} section covers critical aspects of system functionality. Following these guidelines ensures optimal performance and reliability.", styles['Normal']))
        story.append(Spacer(1, 12))
        story.append(Paragraph(f"Best practices for {title.lower()} include regular monitoring, proactive maintenance, and adherence to established protocols. Documentation updates should reflect any configuration changes.", styles['Normal']))

        if i < 20:  # Don't add page break after last page
            story.append(PageBreak())

    doc.build(story)
    print(f"✓ Created 20-page PDF document: {output_path}")
    return output_path

if __name__ == "__main__":
    print("\n" + "="*60)
    print("Creating 20-page test documents...")
    print("="*60 + "\n")

    try:
        word_path = create_20_page_word_doc()
        pdf_path = create_20_page_pdf()

        print("\n" + "="*60)
        print("✓ Successfully created test documents!")
        print("="*60)
        print(f"\nWord Document: {word_path}")
        print(f"PDF Document:  {pdf_path}")
        print(f"\nBoth documents contain 20 pages with unique content on each page.")
        print("Ready for classification and search testing.\n")

    except Exception as e:
        print(f"\n✗ Error creating documents: {e}")
        import traceback
        traceback.print_exc()
