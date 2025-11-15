#!/usr/bin/env python3
"""
Create comprehensive test documents for the AI Document Pipeline.
Generates Word docs, PDFs, and images with various content types.
"""

import os
from pathlib import Path
from datetime import datetime
import json

# Try to import document generation libraries
try:
    from docx import Document
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from PIL import Image, ImageDraw, ImageFont
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False

def create_test_directories():
    """Create test document directories."""
    base_dir = Path("test_documents_comprehensive")
    
    dirs = [
        "word_docs",
        "pdfs", 
        "images",
        "mixed_formats"
    ]
    
    for dir_name in dirs:
        (base_dir / dir_name).mkdir(parents=True, exist_ok=True)
    
    return base_dir

def create_word_documents(base_dir):
    """Create Word document test files."""
    if not DOCX_AVAILABLE:
        print("⚠️  python-docx not available. Install with: pip install python-docx")
        return []
    
    docs_created = []
    word_dir = base_dir / "word_docs"
    
    # 1. Business Contract
    doc = Document()
    doc.add_heading('Software Development Contract', 0)
    doc.add_heading('ACME Technical Consulting', level=1)
    doc.add_paragraph('Contract Number: CTR-2025-001')
    doc.add_paragraph('Date: November 4, 2025')
    doc.add_paragraph('')
    doc.add_paragraph('This contract outlines the software development services to be provided by ACME Technical Consulting for the client\'s enterprise application modernization project.')
    doc.add_paragraph('')
    doc.add_paragraph('Scope of Work:')
    doc.add_paragraph('• Full-stack web application development')
    doc.add_paragraph('• Database design and optimization') 
    doc.add_paragraph('• API development and integration')
    doc.add_paragraph('• Quality assurance and testing')
    doc.add_paragraph('• Code review and documentation')
    doc.add_paragraph('')
    doc.add_paragraph('Total Contract Value: $85,000')
    doc.add_paragraph('Timeline: 6 months')
    
    contract_path = word_dir / "contract_acme_2025.docx"
    doc.save(contract_path)
    docs_created.append(contract_path)
    
    # 2. Meeting Minutes
    doc = Document()
    doc.add_heading('Project Kickoff Meeting Minutes', 0)
    doc.add_paragraph('Date: November 1, 2025')
    doc.add_paragraph('Attendees: John Smith (PM), Sarah Chen (Tech Lead), Mike Johnson (Client)')
    doc.add_paragraph('')
    doc.add_paragraph('Agenda Items Discussed:')
    doc.add_paragraph('1. Project timeline and milestones')
    doc.add_paragraph('2. Technical architecture decisions')
    doc.add_paragraph('3. Resource allocation and team structure')
    doc.add_paragraph('4. Communication protocols')
    doc.add_paragraph('')
    doc.add_paragraph('Action Items:')
    doc.add_paragraph('• Set up development environment by Nov 5')
    doc.add_paragraph('• Finalize database schema by Nov 8')
    doc.add_paragraph('• Schedule weekly standup meetings')
    
    minutes_path = word_dir / "meeting_minutes_kickoff.docx"
    doc.save(minutes_path)
    docs_created.append(minutes_path)
    
    # 3. Technical Specification
    doc = Document()
    doc.add_heading('API Specification Document', 0)
    doc.add_heading('Document Management System', level=1)
    doc.add_paragraph('Version: 1.0')
    doc.add_paragraph('Author: Technical Team')
    doc.add_paragraph('')
    doc.add_paragraph('Overview:')
    doc.add_paragraph('This document outlines the REST API specifications for the document management system, including endpoints for document upload, search, and retrieval.')
    doc.add_paragraph('')
    doc.add_paragraph('Authentication:')
    doc.add_paragraph('All API requests require JWT token authentication.')
    doc.add_paragraph('')
    doc.add_paragraph('Endpoints:')
    doc.add_paragraph('POST /api/documents - Upload new document')
    doc.add_paragraph('GET /api/search - Search documents with filters')
    doc.add_paragraph('GET /api/documents/{id} - Retrieve specific document')
    doc.add_paragraph('DELETE /api/documents/{id} - Delete document')
    
    spec_path = word_dir / "api_specification.docx"
    doc.save(spec_path)
    docs_created.append(spec_path)
    
    return docs_created

def create_pdf_documents(base_dir):
    """Create PDF test files."""
    if not PDF_AVAILABLE:
        print("⚠️  reportlab not available. Install with: pip install reportlab")
        return []
    
    docs_created = []
    pdf_dir = base_dir / "pdfs"
    
    # 1. Invoice PDF
    invoice_path = pdf_dir / "invoice_2025_001.pdf"
    doc = SimpleDocTemplate(str(invoice_path), pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    
    story.append(Paragraph("INVOICE", styles['Title']))
    story.append(Spacer(1, 12))
    story.append(Paragraph("ACME Technical Consulting", styles['Heading1']))
    story.append(Paragraph("123 Tech Street, Silicon Valley, CA 94000", styles['Normal']))
    story.append(Spacer(1, 12))
    story.append(Paragraph("Invoice #: INV-2025-001", styles['Normal']))
    story.append(Paragraph("Date: November 4, 2025", styles['Normal']))
    story.append(Paragraph("Due Date: December 4, 2025", styles['Normal']))
    story.append(Spacer(1, 12))
    story.append(Paragraph("Bill To:", styles['Heading2']))
    story.append(Paragraph("Enterprise Solutions Inc.", styles['Normal']))
    story.append(Paragraph("456 Business Ave, Corporate City, NY 10001", styles['Normal']))
    story.append(Spacer(1, 12))
    story.append(Paragraph("Description of Services:", styles['Heading2']))
    story.append(Paragraph("• Software development consulting - 80 hours @ $150/hour", styles['Normal']))
    story.append(Paragraph("• Code review and quality assurance - 20 hours @ $125/hour", styles['Normal']))
    story.append(Paragraph("• Technical documentation - 10 hours @ $100/hour", styles['Normal']))
    story.append(Spacer(1, 12))
    story.append(Paragraph("Total Amount: $16,500.00", styles['Heading2']))
    
    doc.build(story)
    docs_created.append(invoice_path)
    
    # 2. Financial Report PDF
    report_path = pdf_dir / "financial_report_q3_2025.pdf"
    doc = SimpleDocTemplate(str(report_path), pagesize=letter)
    story = []
    
    story.append(Paragraph("Q3 2025 Financial Report", styles['Title']))
    story.append(Spacer(1, 12))
    story.append(Paragraph("Executive Summary", styles['Heading1']))
    story.append(Paragraph("The third quarter of 2025 showed strong performance across all business units, with revenue increasing 15% compared to Q3 2024.", styles['Normal']))
    story.append(Spacer(1, 12))
    story.append(Paragraph("Revenue Breakdown:", styles['Heading2']))
    story.append(Paragraph("• Consulting Services: $245,000", styles['Normal']))
    story.append(Paragraph("• Software Licensing: $89,000", styles['Normal']))
    story.append(Paragraph("• Training and Support: $34,000", styles['Normal']))
    story.append(Spacer(1, 12))
    story.append(Paragraph("Key Metrics:", styles['Heading2']))
    story.append(Paragraph("• Gross Margin: 68%", styles['Normal']))
    story.append(Paragraph("• Client Retention Rate: 94%", styles['Normal']))
    story.append(Paragraph("• Employee Utilization: 87%", styles['Normal']))
    story.append(Spacer(1, 12))
    story.append(Paragraph("Looking ahead to Q4, we anticipate continued growth with three major client projects scheduled to begin.", styles['Normal']))
    
    doc.build(story)
    docs_created.append(report_path)
    
    return docs_created

def create_image_documents(base_dir):
    """Create image documents with text content."""
    if not PIL_AVAILABLE:
        print("⚠️  Pillow not available. Install with: pip install Pillow")
        return []
    
    docs_created = []
    img_dir = base_dir / "images"
    
    # 1. Business Card Image
    img = Image.new('RGB', (600, 350), color='white')
    draw = ImageDraw.Draw(img)
    
    # Try to use a default font, fallback to basic if not available
    try:
        font_large = ImageFont.truetype("/System/Library/Fonts/Arial.ttf", 24)
        font_medium = ImageFont.truetype("/System/Library/Fonts/Arial.ttf", 18)
        font_small = ImageFont.truetype("/System/Library/Fonts/Arial.ttf", 14)
    except:
        font_large = ImageFont.load_default()
        font_medium = ImageFont.load_default()
        font_small = ImageFont.load_default()
    
    # Draw business card content
    draw.text((50, 50), "ACME Technical Consulting", fill='black', font=font_large)
    draw.text((50, 90), "John Smith", fill='blue', font=font_medium)
    draw.text((50, 120), "Senior Software Engineer", fill='gray', font=font_small)
    draw.text((50, 150), "Email: john.smith@acmetech.com", fill='black', font=font_small)
    draw.text((50, 180), "Phone: (555) 123-4567", fill='black', font=font_small)
    draw.text((50, 210), "Website: www.acmetech.com", fill='black', font=font_small)
    draw.text((50, 250), "Specializing in enterprise software", fill='gray', font=font_small)
    draw.text((50, 270), "development and cloud architecture", fill='gray', font=font_small)
    
    card_path = img_dir / "business_card_john_smith.png"
    img.save(card_path)
    docs_created.append(card_path)
    
    # 2. Handwritten Note Image
    img = Image.new('RGB', (800, 600), color='#f8f8f8')
    draw = ImageDraw.Draw(img)
    
    # Simulate handwritten note
    draw.text((100, 100), "Meeting Notes - November 3, 2025", fill='blue', font=font_medium)
    draw.text((100, 150), "Client Requirements:", fill='black', font=font_small)
    draw.text((120, 180), "- Need document classification system", fill='black', font=font_small)
    draw.text((120, 210), "- OCR capability for scanned documents", fill='black', font=font_small)
    draw.text((120, 240), "- Semantic search functionality", fill='black', font=font_small)
    draw.text((120, 270), "- Integration with existing workflow", fill='black', font=font_small)
    draw.text((100, 320), "Timeline: 3 months", fill='red', font=font_small)
    draw.text((100, 350), "Budget: $50,000", fill='red', font=font_small)
    draw.text((100, 400), "Next Steps:", fill='black', font=font_small)
    draw.text((120, 430), "1. Technical architecture review", fill='black', font=font_small)
    draw.text((120, 460), "2. Prototype development", fill='black', font=font_small)
    draw.text((120, 490), "3. Client demo scheduled for Nov 15", fill='black', font=font_small)
    
    note_path = img_dir / "handwritten_meeting_notes.png"
    img.save(note_path)
    docs_created.append(note_path)
    
    # 3. Scanned Receipt Image
    img = Image.new('RGB', (400, 600), color='white')
    draw = ImageDraw.Draw(img)
    
    # Receipt header
    draw.text((100, 50), "OFFICE DEPOT", fill='black', font=font_medium)
    draw.text((80, 80), "Store #1234", fill='black', font=font_small)
    draw.text((50, 110), "Date: 11/03/2025  Time: 14:30", fill='black', font=font_small)
    
    # Receipt items
    draw.text((50, 150), "HP Printer Paper (5 reams)", fill='black', font=font_small)
    draw.text((300, 150), "$24.99", fill='black', font=font_small)
    draw.text((50, 180), "Black Ink Cartridge", fill='black', font=font_small)
    draw.text((300, 180), "$89.99", fill='black', font=font_small)
    draw.text((50, 210), "Office Supplies", fill='black', font=font_small)
    draw.text((300, 210), "$15.47", fill='black', font=font_small)
    
    # Total
    draw.text((50, 250), "Subtotal:", fill='black', font=font_small)
    draw.text((300, 250), "$130.45", fill='black', font=font_small)
    draw.text((50, 280), "Tax:", fill='black', font=font_small)
    draw.text((300, 280), "$10.44", fill='black', font=font_small)
    draw.text((50, 310), "TOTAL:", fill='black', font=font_medium)
    draw.text((300, 310), "$140.89", fill='black', font=font_medium)
    
    draw.text((50, 350), "Payment Method: Credit Card", fill='black', font=font_small)
    draw.text((50, 380), "Card ending in: ****1234", fill='black', font=font_small)
    
    receipt_path = img_dir / "receipt_office_supplies.png"
    img.save(receipt_path)
    docs_created.append(receipt_path)
    
    return docs_created

def create_summary_file(base_dir, all_docs):
    """Create a summary file listing all test documents."""
    summary = {
        "created_date": datetime.now().isoformat(),
        "total_documents": len(all_docs),
        "document_types": {
            "word_docs": len([d for d in all_docs if d.suffix == '.docx']),
            "pdfs": len([d for d in all_docs if d.suffix == '.pdf']),
            "images": len([d for d in all_docs if d.suffix in ['.png', '.jpg', '.jpeg']])
        },
        "files": [str(doc.relative_to(base_dir)) for doc in all_docs],
        "usage_instructions": {
            "copy_to_input": f"cp test_documents_comprehensive/*/* documents/input/",
            "process_documents": "python3 -m src.cli process documents/input/",
            "test_search": "curl -X GET 'http://localhost:8000/api/search?q=ACME&mode=semantic'"
        }
    }
    
    summary_path = base_dir / "test_documents_summary.json"
    with open(summary_path, 'w') as f:
        json.dump(summary, f, indent=2)
    
    return summary_path

def main():
    print("🚀 Creating comprehensive test documents...")
    print("=" * 50)
    
    # Create directories
    base_dir = create_test_directories()
    print(f"📁 Created base directory: {base_dir}")
    
    all_docs = []
    
    # Create Word documents
    print("\n📄 Creating Word documents...")
    word_docs = create_word_documents(base_dir)
    all_docs.extend(word_docs)
    print(f"✅ Created {len(word_docs)} Word documents")
    
    # Create PDF documents  
    print("\n📋 Creating PDF documents...")
    pdf_docs = create_pdf_documents(base_dir)
    all_docs.extend(pdf_docs)
    print(f"✅ Created {len(pdf_docs)} PDF documents")
    
    # Create image documents
    print("\n🖼️  Creating image documents...")
    image_docs = create_image_documents(base_dir)
    all_docs.extend(image_docs)
    print(f"✅ Created {len(image_docs)} image documents")
    
    # Create summary
    summary_path = create_summary_file(base_dir, all_docs)
    print(f"\n📊 Created summary file: {summary_path}")
    
    print(f"\n🎉 Successfully created {len(all_docs)} test documents!")
    print("\n📖 Next steps:")
    print(f"1. Copy documents to input folder:")
    print(f"   cp {base_dir}/*/* documents/input/")
    print("\n2. Process documents:")
    print("   python3 -m src.cli process documents/input/")
    print("\n3. Test search:")
    print("   curl -X GET 'http://localhost:8000/api/search?q=ACME&mode=semantic'")
    
    print(f"\n📂 Documents created in: {base_dir.absolute()}")

if __name__ == "__main__":
    main()