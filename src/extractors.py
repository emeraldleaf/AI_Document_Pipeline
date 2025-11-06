"""Document content and metadata extraction for various file formats."""

import mimetypes
from abc import ABC, abstractmethod
from pathlib import Path
from typing import Dict, Any, List, Optional
from datetime import datetime

import pdfplumber
from docx import Document as DocxDocument
import openpyxl
import pandas as pd
from loguru import logger
from pdf2image import convert_from_path
from PIL import Image

# Import the image processor
try:
    from image_processor import ImageProcessor
except ImportError:
    # Try relative import as fallback
    try:
        from .image_processor import ImageProcessor
    except ImportError:
        ImageProcessor = None


class DocumentMetadata:
    """Container for document metadata."""

    def __init__(self, file_path: Path, **kwargs):
        """Initialize metadata with file path and optional properties."""
        self.file_path = file_path
        self.file_name = kwargs.get('file_name', file_path.name)
        self.file_type = kwargs.get('file_type', 'unknown')
        self.file_size = kwargs.get('file_size', 0)
        self.created_date = kwargs.get('created_date')
        self.modified_date = kwargs.get('modified_date')
        self.page_count = kwargs.get('page_count')
        self.author = kwargs.get('author')
        self.title = kwargs.get('title')

    @classmethod
    def from_file_stats(cls, file_path: Path) -> 'DocumentMetadata':
        """Create metadata from file system statistics."""
        stat = file_path.stat()
        mime_type, _ = mimetypes.guess_type(file_path)

        return cls(
            file_path=file_path,
            file_name=file_path.name,
            file_type=mime_type or "unknown",
            file_size=stat.st_size,
            created_date=datetime.fromtimestamp(stat.st_ctime),
            modified_date=datetime.fromtimestamp(stat.st_mtime),
        )

    def to_dict(self) -> Dict[str, Any]:
        """Convert metadata to dictionary."""
        return {
            "file_path": str(self.file_path),
            "file_name": self.file_name,
            "file_type": self.file_type,
            "file_size": self.file_size,
            "created_date": self.created_date.isoformat() if self.created_date else None,
            "modified_date": self.modified_date.isoformat() if self.modified_date else None,
            "page_count": self.page_count,
            "author": self.author,
            "title": self.title,
        }


class ExtractedContent:
    """Container for extracted document content."""

    def __init__(self, text: str, metadata: DocumentMetadata, pages: Optional[List[str]] = None):
        self.text = text
        self.metadata = metadata
        self.pages = pages or []

    def get_summary(self, max_length: int = 500) -> str:
        """Get a truncated summary of the content."""
        if len(self.text) <= max_length:
            return self.text
        return self.text[:max_length] + "..."


class BaseExtractor(ABC):
    """Base class for document extractors."""

    @abstractmethod
    def can_extract(self, file_path: Path) -> bool:
        """Check if this extractor can handle the given file."""
        pass

    @abstractmethod
    def extract(self, file_path: Path) -> ExtractedContent:
        """Extract content and metadata from the document."""
        pass

    def _get_basic_metadata(self, file_path: Path) -> DocumentMetadata:
        """Extract basic file system metadata."""
        return DocumentMetadata.from_file_stats(file_path)


class PDFExtractor(BaseExtractor):
    """Extractor for PDF documents with OCR support for image-only PDFs."""

    def __init__(self):
        """Initialize PDFExtractor with optional OCR support."""
        self.ocr_processor = None
        try:
            self.ocr_processor = ImageProcessor()
            logger.info("PDF extractor initialized with OCR support")
        except Exception as e:
            logger.warning(f"OCR not available for PDF extraction: {e}")

    def can_extract(self, file_path: Path) -> bool:
        """Check if file is a PDF."""
        return file_path.suffix.lower() == ".pdf"

    def _has_extractable_text(self, pdf) -> bool:
        """Check if PDF has extractable text content."""
        for page in pdf.pages[:3]:  # Check first 3 pages
            text = page.extract_text()
            if text and text.strip():
                # Check if it's meaningful text (not just whitespace/symbols)
                words = text.split()
                if len(words) > 5:  # Arbitrary threshold
                    return True
        return False

    def _extract_with_ocr(self, file_path: Path, metadata: DocumentMetadata) -> ExtractedContent:
        """Extract text from image-based PDF using OCR."""
        if not self.ocr_processor:
            raise RuntimeError("OCR processor not available")
        
        logger.info(f"Using OCR for image-based PDF: {file_path.name}")
        
        try:
            # Convert PDF pages to images
            images = convert_from_path(file_path, dpi=300)
            metadata.page_count = len(images)
            
            text_content = []
            page_texts = []
            total_confidence = 0
            
            for page_num, image in enumerate(images, 1):
                logger.debug(f"Processing page {page_num} with OCR")
                
                # Extract text from the page image
                ocr_result = self.ocr_processor.extract_text_from_pdf_image(image)
                page_text = ocr_result['text']
                
                if page_text.strip():
                    page_texts.append(page_text)
                    text_content.append(f"[Page {page_num}]\n{page_text}")
                    total_confidence += ocr_result['confidence']
                else:
                    page_texts.append("")
                    text_content.append(f"[Page {page_num}]\n[No text found]")
            
            # Calculate average confidence
            avg_confidence = total_confidence / len(images) if images else 0
            
            full_text = "\n\n".join(text_content)
            
            # Add OCR metadata
            if hasattr(metadata, 'ocr_confidence'):
                metadata.ocr_confidence = avg_confidence
            
            logger.success(f"OCR extracted text from {len(images)} pages (avg confidence: {avg_confidence:.1f}%)")
            
            return ExtractedContent(text=full_text, metadata=metadata, pages=page_texts)
            
        except Exception as e:
            logger.error(f"OCR extraction failed for {file_path.name}: {e}")
            raise

    def extract(self, file_path: Path) -> ExtractedContent:
        """Extract text and metadata from PDF with fallback to OCR."""
        logger.info(f"Extracting PDF: {file_path.name}")

        metadata = self._get_basic_metadata(file_path)
        text_content = []
        page_texts = []

        try:
            with pdfplumber.open(file_path) as pdf:
                metadata.page_count = len(pdf.pages)

                # Extract PDF metadata if available
                if pdf.metadata:
                    metadata.author = pdf.metadata.get("Author")
                    metadata.title = pdf.metadata.get("Title")

                # Check if PDF has extractable text
                if self._has_extractable_text(pdf):
                    # Extract text normally
                    for page_num, page in enumerate(pdf.pages, 1):
                        page_text = page.extract_text() or ""
                        page_texts.append(page_text)
                        text_content.append(f"[Page {page_num}]\n{page_text}")

                    full_text = "\n\n".join(text_content)
                    logger.success(f"Extracted {len(page_texts)} pages from {file_path.name}")

                    return ExtractedContent(text=full_text, metadata=metadata, pages=page_texts)
                
                else:
                    # PDF appears to be image-based, use OCR
                    if self.ocr_processor:
                        logger.info(f"PDF appears to be image-based, switching to OCR: {file_path.name}")
                        return self._extract_with_ocr(file_path, metadata)
                    else:
                        logger.warning(f"Image-based PDF detected but OCR not available: {file_path.name}")
                        # Return empty content rather than failing
                        return ExtractedContent(text="[Image-based PDF - OCR not available]", metadata=metadata)

        except Exception as e:
            logger.error(f"Error extracting PDF {file_path.name}: {e}")
            raise


class ImageExtractor(BaseExtractor):
    """Extractor for image documents using OCR."""

    def __init__(self):
        """Initialize ImageExtractor with OCR processor."""
        try:
            self.ocr_processor = ImageProcessor()
            logger.info("Image extractor initialized with OCR support")
        except Exception as e:
            logger.error(f"Failed to initialize OCR for image extraction: {e}")
            raise

    def can_extract(self, file_path: Path) -> bool:
        """Check if file is a supported image format."""
        return self.ocr_processor.is_supported(str(file_path))

    def extract(self, file_path: Path) -> ExtractedContent:
        """Extract text from image using OCR."""
        logger.info(f"Extracting image: {file_path.name}")

        metadata = self._get_basic_metadata(file_path)

        try:
            # Use OCR to extract text
            ocr_result = self.ocr_processor.extract_text(str(file_path))
            
            text = ocr_result['text']
            confidence = ocr_result['confidence']
            
            # Set page count to 1 for images
            metadata.page_count = 1
            
            # Add OCR-specific metadata
            if hasattr(metadata, 'ocr_confidence'):
                metadata.ocr_confidence = confidence
            
            # Add image-specific metadata
            if hasattr(metadata, 'image_size'):
                metadata.image_size = ocr_result['image_size']
            if hasattr(metadata, 'image_format'):
                metadata.image_format = ocr_result['format']

            if text.strip():
                logger.success(f"Extracted text from {file_path.name} (confidence: {confidence:.1f}%)")
            else:
                logger.warning(f"No text found in {file_path.name}")
                text = "[No text detected in image]"

            return ExtractedContent(text=text, metadata=metadata)

        except Exception as e:
            logger.error(f"Error extracting image {file_path.name}: {e}")
            raise


class DOCXExtractor(BaseExtractor):
    """Extractor for Word documents."""

    def can_extract(self, file_path: Path) -> bool:
        """Check if file is a DOCX."""
        return file_path.suffix.lower() in [".docx", ".doc"]

    def extract(self, file_path: Path) -> ExtractedContent:
        """Extract text and metadata from DOCX."""
        logger.info(f"Extracting DOCX: {file_path.name}")

        metadata = self._get_basic_metadata(file_path)

        try:
            doc = DocxDocument(file_path)

            # Extract metadata
            core_properties = doc.core_properties
            metadata.author = core_properties.author
            metadata.title = core_properties.title
            metadata.page_count = len(doc.sections)

            # Extract text from paragraphs
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            full_text = "\n\n".join(paragraphs)

            # Extract text from tables
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    tables_text.append(row_text)

            if tables_text:
                full_text += "\n\n[Tables]\n" + "\n".join(tables_text)

            logger.success(f"Extracted {len(paragraphs)} paragraphs from {file_path.name}")

            return ExtractedContent(text=full_text, metadata=metadata)

        except Exception as e:
            logger.error(f"Error extracting DOCX {file_path.name}: {e}")
            raise


class ExcelExtractor(BaseExtractor):
    """Extractor for Excel spreadsheets."""

    def can_extract(self, file_path: Path) -> bool:
        """Check if file is an Excel file."""
        return file_path.suffix.lower() in [".xlsx", ".xls", ".xlsm"]

    def extract(self, file_path: Path) -> ExtractedContent:
        """Extract text and metadata from Excel."""
        logger.info(f"Extracting Excel: {file_path.name}")

        metadata = self._get_basic_metadata(file_path)

        try:
            # Load workbook
            wb = openpyxl.load_workbook(file_path, data_only=True)
            metadata.page_count = len(wb.sheetnames)

            # Extract metadata
            if wb.properties:
                metadata.author = wb.properties.creator
                metadata.title = wb.properties.title

            # Extract content from all sheets
            sheets_content = []
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                sheets_content.append(f"[Sheet: {sheet_name}]")

                # Convert sheet to pandas DataFrame for easier text extraction
                data = []
                for row in ws.iter_rows(values_only=True):
                    data.append([str(cell) if cell is not None else "" for cell in row])

                if data:
                    df = pd.DataFrame(data)
                    sheet_text = df.to_string(index=False, header=False)
                    sheets_content.append(sheet_text)

            full_text = "\n\n".join(sheets_content)
            logger.success(f"Extracted {len(wb.sheetnames)} sheets from {file_path.name}")

            return ExtractedContent(text=full_text, metadata=metadata)

        except Exception as e:
            logger.error(f"Error extracting Excel {file_path.name}: {e}")
            raise


class TextExtractor(BaseExtractor):
    """Extractor for plain text files."""

    def can_extract(self, file_path: Path) -> bool:
        """Check if file is a text file."""
        return file_path.suffix.lower() in [".txt", ".md", ".csv", ".json", ".xml"]

    def extract(self, file_path: Path) -> ExtractedContent:
        """Extract text from plain text file."""
        logger.info(f"Extracting text file: {file_path.name}")

        metadata = self._get_basic_metadata(file_path)

        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()

            metadata.page_count = 1
            logger.success(f"Extracted text from {file_path.name}")

            return ExtractedContent(text=text, metadata=metadata)

        except Exception as e:
            logger.error(f"Error extracting text file {file_path.name}: {e}")
            raise


class ExtractionService:
    """Service to coordinate document extraction across different formats."""

    def __init__(self):
        self.extractors: List[BaseExtractor] = []
        
        # Initialize PDF extractor
        try:
            self.extractors.append(PDFExtractor())
        except Exception as e:
            logger.warning(f"PDF extractor initialization failed: {e}")
            # Fall back to basic PDF extractor without OCR
            self.extractors.append(self._create_basic_pdf_extractor())
        
        # Initialize Image extractor
        try:
            self.extractors.append(ImageExtractor())
        except Exception as e:
            logger.warning(f"Image extractor initialization failed (OCR not available): {e}")
        
        # Add other extractors
        self.extractors.extend([
            DOCXExtractor(),
            ExcelExtractor(),
            TextExtractor(),
        ])

    def _create_basic_pdf_extractor(self):
        """Create a basic PDF extractor without OCR capabilities."""
        class BasicPDFExtractor(BaseExtractor):
            def can_extract(self, file_path: Path) -> bool:
                return file_path.suffix.lower() == ".pdf"
            
            def extract(self, file_path: Path) -> ExtractedContent:
                logger.info(f"Extracting PDF (basic mode): {file_path.name}")
                metadata = self._get_basic_metadata(file_path)
                text_content = []
                page_texts = []

                try:
                    with pdfplumber.open(file_path) as pdf:
                        metadata.page_count = len(pdf.pages)
                        if pdf.metadata:
                            metadata.author = pdf.metadata.get("Author")
                            metadata.title = pdf.metadata.get("Title")

                        for page_num, page in enumerate(pdf.pages, 1):
                            page_text = page.extract_text() or ""
                            page_texts.append(page_text)
                            text_content.append(f"[Page {page_num}]\n{page_text}")

                    full_text = "\n\n".join(text_content)
                    return ExtractedContent(text=full_text, metadata=metadata, pages=page_texts)
                except Exception as e:
                    logger.error(f"Error extracting PDF {file_path.name}: {e}")
                    raise
        
        return BasicPDFExtractor()

    def extract(self, file_path: Path) -> Optional[ExtractedContent]:
        """Extract content from a document using the appropriate extractor."""
        if not file_path.exists():
            logger.error(f"File not found: {file_path}")
            return None

        for extractor in self.extractors:
            if extractor.can_extract(file_path):
                try:
                    return extractor.extract(file_path)
                except Exception as e:
                    logger.error(f"Failed to extract {file_path.name}: {e}")
                    return None

        logger.warning(f"No extractor found for {file_path.name}")
        return None

    def extract_batch(self, file_paths: List[Path]) -> List[ExtractedContent]:
        """Extract content from multiple documents."""
        results = []
        for file_path in file_paths:
            content = self.extract(file_path)
            if content:
                results.append(content)
        return results
